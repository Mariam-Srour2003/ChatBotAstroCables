"""
build_vectorstore.py
One-time setup script: load ALL Astro documents and build the FAISS index.

Run once before starting the chatbot:
    python build_vectorstore.py

Supported file types:
    .txt   cable spec text files
    .docx  company documents
    .xlsx  cable electrical / mechanical data tables
    .csv   additional cable specs
"""
import csv
import glob
import os
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

from app.config import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    DOCS_DIR,
    EMBEDDING_MODEL,
    VECTORSTORE_PATH,
)

# ── Loaders ───────────────────────────────────────────────────────────────────

def load_txt(directory: str) -> list[Document]:
    docs = []
    for path in glob.glob(f"{directory}/**/*.txt", recursive=True):
        try:
            docs.extend(TextLoader(path, encoding="utf-8").load())
        except Exception as e:
            print(f"  [skip] {path}: {e}")
    return docs


def load_docx(directory: str) -> list[Document]:
    from docx import Document as DocxDoc
    docs = []
    for path in glob.glob(f"{directory}/**/*.docx", recursive=True):
        try:
            doc  = DocxDoc(path)
            # Paragraphs
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Table cells
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        c.text.strip() for c in row.cells if c.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            text = "\n".join(parts)
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": path, "type": "docx"},
                ))
        except Exception as e:
            print(f"  [skip] {path}: {e}")
    return docs


def load_xlsx(directory: str) -> list[Document]:
    import openpyxl
    docs = []
    for path in glob.glob(f"{directory}/**/*.xlsx", recursive=True):
        try:
            wb         = openpyxl.load_workbook(path, data_only=True)
            cable_name = Path(path).parent.name
            for sheet in wb.sheetnames:
                ws   = wb[sheet]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells    = [str(c) if c is not None else "" for c in row]
                    row_text = " | ".join(cells)
                    if row_text.replace("|", "").strip():
                        rows.append(row_text)
                if rows:
                    text = (
                        f"Cable Type: {cable_name}\n"
                        f"Data Sheet: {sheet}\n"
                        + "\n".join(rows)
                    )
                    docs.append(Document(
                        page_content=text,
                        metadata={"source": path, "cable": cable_name, "sheet": sheet, "type": "xlsx"},
                    ))
        except Exception as e:
            print(f"  [skip] {path}: {e}")
    return docs


def load_csv(directory: str) -> list[Document]:
    docs = []
    for path in glob.glob(f"{directory}/**/*.csv", recursive=True):
        try:
            with open(path, "r", encoding="utf-8-sig", errors="ignore") as f:
                rows = [
                    " | ".join(cell.strip() for cell in row)
                    for row in csv.reader(f)
                    if any(cell.strip() for cell in row)
                ]
            if rows:
                cable_name = Path(path).stem
                text       = f"Cable Type: {cable_name}\n" + "\n".join(rows)
                docs.append(Document(
                    page_content=text,
                    metadata={"source": path, "cable": cable_name, "type": "csv"},
                ))
        except Exception as e:
            print(f"  [skip] {path}: {e}")
    return docs


# ── Main ──────────────────────────────────────────────────────────────────────

def build() -> None:
    print("=" * 60)
    print("  Astro Power Cables — Building Vectorstore")
    print("=" * 60)

    loaders = [
        (".txt  files", load_txt),
        (".docx files", load_docx),
        (".xlsx files", load_xlsx),
        (".csv  files", load_csv),
    ]

    all_docs: list[Document] = []
    print(f"\nReading from: {DOCS_DIR}\n")
    for label, fn in loaders:
        batch = fn(DOCS_DIR)
        print(f"  {label}: {len(batch):>4} document(s)")
        all_docs.extend(batch)

    if not all_docs:
        print("\nNo documents found — check DOCS_DIR in app/config.py.")
        return

    print(f"\nTotal raw documents : {len(all_docs)}")

    print("Splitting into chunks...")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_documents(all_docs)
    print(f"Total chunks        : {len(chunks)}")

    print("\nBuilding FAISS index (may take a minute)...")
    embedding = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )
    db = FAISS.from_documents(chunks, embedding)

    os.makedirs(VECTORSTORE_PATH, exist_ok=True)
    db.save_local(VECTORSTORE_PATH)
    print(f"\nVectorstore saved to: {VECTORSTORE_PATH}")
    print("Done. Run  python main.py  or  uvicorn app.api:app  to start.\n")


if __name__ == "__main__":
    build()
